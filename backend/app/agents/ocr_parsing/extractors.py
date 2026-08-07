"""Format detection and text extraction. Native parsing for digital PDF/DOCX,
Tesseract for scanned pages."""
import io
from typing import Literal

import pdfplumber
import pytesseract
from docx import Document as DocxDocument
from pdf2image import convert_from_bytes

SourceFormat = Literal["pdf", "docx", "scanned"]

# Below this ratio of extractable characters per page, treat the PDF as scanned.
DIGITAL_TEXT_THRESHOLD = 100


def _is_digital(text: str) -> bool:
    """Above this ratio of extractable characters per page, the PDF is digital
    text rather than a scanned image."""
    pages = max(text.count("\f"), 1)
    return len(text.strip()) / pages >= DIGITAL_TEXT_THRESHOLD


def detect_and_extract(data: bytes, filename: str) -> tuple[SourceFormat, str, float]:
    """Returns (source_format, raw_text, confidence)."""
    lower = filename.lower()
    if lower.endswith(".docx"):
        return "docx", _extract_docx(data), 1.0
    if lower.endswith(".pdf"):
        text = _extract_pdf_native(data)
        if _is_digital(text):
            return "pdf", text, 1.0
        ocr_text, confidence = _extract_scanned(data)
        return "scanned", ocr_text, confidence
    ocr_text, confidence = _extract_scanned(data)
    return "scanned", ocr_text, confidence


def detect_format(data: bytes, filename: str) -> SourceFormat:
    """Coarse format label computed at upload time, matching what the OCR agent
    will later store in the payload. Runs no OCR - a PDF is classed as scanned
    only when native text extraction is thin, which mirrors detect_and_extract."""
    lower = filename.lower()
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".pdf"):
        try:
            return "pdf" if _is_digital(_extract_pdf_native(data)) else "scanned"
        except Exception:
            # Unreadable PDF. Extension says pdf; the OCR agent will surface
            # the actual parse failure rather than mislabeling it scanned.
            return "pdf"
    return "scanned"


def _extract_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_pdf_native(data: bytes) -> str:
    out = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\f".join(out)


def _extract_scanned(data: bytes) -> tuple[str, float]:
    images = convert_from_bytes(data, dpi=300)
    texts, confidences = [], []
    for image in images:
        texts.append(pytesseract.image_to_string(image))
        report = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        page_scores = [int(c) for c in report["conf"] if c not in ("-1", -1)]
        if page_scores:
            confidences.append(sum(page_scores) / len(page_scores) / 100)
    confidence = sum(confidences) / len(confidences) if confidences else 0.0
    return "\f".join(texts), round(confidence, 3)
